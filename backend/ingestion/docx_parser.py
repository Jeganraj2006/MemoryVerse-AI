"""DOCX text extraction."""
from __future__ import annotations

import asyncio


async def extract_docx_text(docx_path: str) -> dict:
    return await asyncio.to_thread(_extract_docx_sync, docx_path)


def _extract_docx_sync(docx_path: str) -> dict:
    from docx import Document

    document = Document(docx_path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                paragraphs.append(" | ".join(values))
    text = "\n".join(paragraphs)
    return {
        "text": text,
        "method": "python-docx",
        "page_count": 1,
        "pages": [{"page_number": 1, "text": text}],
    }
